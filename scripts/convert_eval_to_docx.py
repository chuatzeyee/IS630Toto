from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

add_heading('IS630 Group Project Ideas — Comprehensive Evaluation', level=0)

add_heading('Evaluation Framework', level=1)
add_para('Each idea is assessed against the IS630 Project Requirements (30% of final grade) on seven criteria:')
add_table(
    ['Criterion', 'What It Measures', 'Weight'],
    [
        ['Course Alignment', 'Fit with frequentist/Bayesian statistical techniques', 'Critical'],
        ['Data Availability', 'Downloadable, clean, large enough for analysis?', 'Critical'],
        ['Novelty', 'New perspective, counter-intuitive hypothesis, original angle?', 'High'],
        ['Statistical Depth', 'Can 6 people each apply distinct methods?', 'High'],
        ['"So What?" Factor', 'Does the conclusion matter to a real audience?', 'Medium'],
        ['Presentation Appeal', 'Engage a non-specialist audience in 15 minutes?', 'Medium'],
        ['Risk Profile', 'Probability of boring/null results or data roadblocks?', 'Medium'],
    ]
)

add_heading('Idea 1: Singapore Weather vs Tourism', level=1)
add_para('Proposed by: Teammate', bold=True)
add_para('Premise: Test whether Singapore\'s climate comfort (temperature, humidity) correlates with tourist arrivals by country of origin.')
add_heading('Datasets', level=2)
add_table(
    ['Dataset', 'Source', 'Size'],
    [
        ['Tourist arrivals by country (monthly)', 'data.gov.sg', '13,700 rows x 4 cols (1978–2015)'],
        ['Surface air temperature (monthly)', 'data.gov.sg', '~526 rows x 2 cols (1982–2025)'],
        ['Relative humidity (monthly)', 'data.gov.sg', '~526 rows x 2 cols (1982–2025)'],
    ]
)
add_heading('Difficulties & Risks', level=2)
for d in [
    'Near-zero signal in the data. Singapore is equatorial. Monthly mean temperature ranges from ~26.0°C to ~28.5°C year-round — a 2.5°C band. Finding a statistically significant relationship is extremely unlikely.',
    'Confounders dwarf the weather signal. Tourism is driven by school holidays, festivals, airline pricing, visa policies, economic conditions, exchange rates — weather is not even in the top 10 factors.',
    'Data ends in 2015. A decade old. The professor may question relevance.',
    'The "by country" angle requires each source country\'s climate profile — which was not identified as a dataset.',
    'Limited statistical variety for 6 people with only monthly data across three tiny datasets.',
]:
    doc.add_paragraph(d, style='List Number')
add_heading('Verdict: 3/10', level=2)
add_para('The fundamental problem is physics: Singapore\'s weather barely varies. You would conclude "weather has no significant effect on tourism in an equatorial city" — obvious without statistics.', italic=True)

add_heading('Idea 2: Telecom Customer Churn', level=1)
add_para('Proposed by: Teammate', bold=True)
add_para('Premise: Examine whether different customer types exhibit different churn behaviour in telecommunications.')
add_heading('Difficulties & Risks', level=2)
for d in [
    'The single most overused project topic in data science. 10,000+ Kaggle notebooks, thousands of Medium articles.',
    'Novelty is nearly impossible. The requirements explicitly ask for "Context and Novelty."',
    'Classification-oriented, not statistics-oriented. IS630 is about hypothesis testing and CIs, not ML prediction.',
    'The IBM dataset is synthetic — no real company, no real business context.',
    'The proposer\'s angle ("different customer types") is what every churn analysis does.',
]:
    doc.add_paragraph(d, style='List Number')
add_heading('Verdict: 2/10', level=2)
add_para('Safest way to get an average grade. Submitting the most overused dataset signals neither innovation nor deep thinking.', italic=True)

add_heading('Idea 3: ESG vs Stock Returns', level=1)
add_para('Proposed by: Teammate', bold=True)
add_para('Premise: Test whether higher ESG scores deliver better stock returns. Break down E, S, G. Compare volatility.')
add_heading('Difficulties & Risks', level=2)
for d in [
    'The ESG dataset is tiny — 503 rows, one per company. Cross-sectional snapshot, not time series.',
    'Data merge is non-trivial. ESG ratings are annual; stock returns are daily.',
    'Survivorship bias — companies removed from S&P 500 aren\'t in the dataset.',
    'Extensive prior research (2,200+ studies in Friede et al. 2015 meta-analysis).',
    'Likely null result. Hard to make compelling 15-minute presentation about "no effect."',
    'Confounders require multivariate regression beyond Sessions 1–5.',
]:
    doc.add_paragraph(d, style='List Number')
add_heading('Verdict: 5/10', level=2)
add_para('Good statistical depth but tiny dataset, complex merge, survivorship bias, and vast prior literature.', italic=True)

add_heading('Idea 4: Spotify — Does Sadder Music Get More Plays? [RECOMMENDED]', level=1)
add_para('Proposed by: Teammate', bold=True)
add_para('Premise: Test whether lower-valence songs get more streams. Segment by genre. Check temporal trends.')
add_heading('Datasets', level=2)
add_table(
    ['Dataset', 'Source', 'Size'],
    [['Spotify Tracks', 'HuggingFace / Kaggle', '114,000 rows x 20 cols']]
)
add_heading('Difficulties (Manageable)', level=2)
for d in [
    'Valence is not sadness — it measures sonic positivity, not lyrics. Acknowledge as limitation.',
    '"Popularity" is not stream count — Spotify\'s metric is recency-weighted.',
    'Confounders (playlist placement, artist fame) will produce low R-squared.',
    'Selection bias — 114K tracks from 100M+ catalog.',
    'Genre segmentation (125 genres) may reduce per-genre sample sizes.',
]:
    doc.add_paragraph(d, style='List Number')
add_heading('Why It Works Despite Difficulties', level=2)
for s in [
    '114,000 rows with 16 numeric features — largest, cleanest dataset among all proposals.',
    'Multiple testable hypotheses: valence vs popularity, genre effects, temporal trends, non-linear effects, multi-feature analysis, explicit content.',
    'Natural 6-way work split — each member owns a distinct hypothesis.',
    'Strongest presentation potential — real song examples engage any audience.',
    'Counter-intuitive premise — perfect for the rubric\'s "investigate counter-intuitive results" tip.',
    'Fits both frequentist AND Bayesian approaches naturally.',
]:
    doc.add_paragraph(s, style='List Bullet')
add_heading('Statistical Methods Fit', level=2)
add_table(
    ['Method', 'Applicability', 'Fit'],
    [
        ['Two-sample t-test', 'Low vs high valence popularity', 'Excellent'],
        ['ANOVA / Kruskal-Wallis', 'Across genres, valence quartiles', 'Excellent'],
        ['Confidence intervals', 'Mean popularity difference', 'Excellent'],
        ['Chi-squared', 'Valence independent of genre?', 'Excellent'],
        ['Regression', 'Popularity ~ valence + features', 'Excellent'],
        ['Bayesian comparison', 'Prior (happy=popular) vs posterior', 'Good'],
        ['Non-parametric tests', 'If popularity is skewed', 'Excellent'],
        ['Normality tests', 'Shapiro-Wilk by group', 'Excellent'],
    ]
)
add_heading('Verdict: 8.5/10 — THE BEST CHOICE', level=2)
add_para('Large clean dataset, counter-intuitive premise, every course method applicable, excellent presentation appeal, natural 6-way work split.', italic=True)

add_heading('Idea 5: Speed Dating', level=1)
add_para('Proposed by: Teammate', bold=True)
add_para('Premise: Analyze what makes someone go on a second date. Self-perception vs others\' perception.')
add_heading('Difficulties & Risks', level=2)
for d in [
    'Severely missing data — some columns 65%+ missing.',
    'Small, culturally narrow — 552 participants from Columbia University, NYC, 2002–2004.',
    'Famous, well-analyzed dataset (Fisman & Iyengar, 2006). Novelty is hard.',
    'Proposer called it "unserious" — risky framing for 30% of final grade.',
    '195 columns create analysis paralysis and multiple comparisons risk.',
    'Class imbalance: 83.5% non-matches.',
]:
    doc.add_paragraph(d, style='List Number')
add_heading('Verdict: 5/10', level=2)
add_para('Fun but missing data burden, cultural narrowness, and "unserious" positioning hurt.', italic=True)

add_heading('Idea 6: Diversification — Always Good?', level=1)
add_para('Proposed by: Teammate', bold=True)
add_para('Premise: Test whether diversification always leads to higher alpha.')
add_heading('Difficulties & Risks', level=2)
for d in [
    'No dataset was identified by the proposer.',
    'This is finance theory (Modern Portfolio Theory, 1952), not statistical analysis.',
    '"Alpha" requires CAPM/Fama-French factor models — specialized finance, not IS630.',
    'Unfocused scope — three distinct questions crammed into one project.',
    'The "overlapping assets" point is trivial — one correlation matrix, one paragraph.',
]:
    doc.add_paragraph(d, style='List Number')
add_heading('Verdict: 2/10', level=2)
add_para('No dataset, no testable hypothesis, no course alignment.', italic=True)

add_heading('Idea 7: Dividends = Higher Growth?', level=1)
add_para('Proposed by: Teammate', bold=True)
add_para('Premise: Compare total returns of dividend-paying vs non-dividend-paying S&P 500 stocks.')
add_heading('Difficulties & Risks', level=2)
for d in [
    'Well-trodden territory — Vanguard, Morningstar, dozens of academics have published extensively.',
    'One main test (two-sample t-test). Too thin for 6 people.',
    'Time period cherry-picking determines the conclusion.',
    'Survivorship bias in S&P 500 datasets.',
    'Dividend classification data not identified.',
]:
    doc.add_paragraph(d, style='List Number')
add_heading('Verdict: 3/10', level=2)
add_para('Too thin for a 6-person team. One main hypothesis test, well-explored territory.', italic=True)

add_heading('Idea 8: Singapore Pools — Lucky Outlets?', level=1)
add_para('Proposed by: Teammate', bold=True)
add_para('Premise: Test whether some TOTO outlets are statistically "luckier," controlling for volume and demographics.')
add_heading('Difficulties & Risks', level=2)
for d in [
    'Data scraping required — Singapore Pools uses dynamic web loading, no CSV download.',
    'Very small dataset — ~300 outlets, many with zero or one wins.',
    '"Tickets sold per outlet" is not public — critical confounding variable unavailable.',
    'The answer is statistically predetermined — lottery outcomes are random.',
    'Rare event statistics (low-count Poisson) require careful handling.',
]:
    doc.add_paragraph(d, style='List Number')
add_heading('Why It Partially Works', level=2)
for s in [
    'Highest novelty — no IS630 team has likely analyzed this before.',
    'Perfect local relevance — debunking a common Singaporean superstition.',
    'Geospatial visualization potential.',
    'Poisson distribution fit — direct Session 3 application.',
    'Bayesian angle — prior vs posterior on "lucky outlets."',
]:
    doc.add_paragraph(s, style='List Bullet')
add_heading('Verdict: 6/10', level=2)
add_para('Most novel, but data scraping and sparsity are serious risks.', italic=True)

add_heading('Comparative Ranking', level=1)
add_table(
    ['Criterion', 'Weather', 'Churn', 'ESG', 'Spotify', 'Dating', 'Diversif.', 'Dividends', 'SG Pools'],
    [
        ['Course Alignment', '4', '3', '6', '9', '6', '2', '4', '7'],
        ['Data Availability', '5', '8', '5', '10', '5', '1', '4', '3'],
        ['Novelty', '4', '1', '4', '7', '3', '3', '2', '10'],
        ['Statistical Depth', '3', '3', '6', '9', '6', '2', '3', '5'],
        ['"So What?"', '3', '4', '5', '7', '3', '4', '4', '9'],
        ['Presentation', '4', '3', '5', '9', '8', '3', '3', '8'],
        ['Low Risk', '3', '6', '4', '7', '5', '2', '4', '4'],
        ['TOTAL (/70)', '26', '28', '35', '58', '36', '17', '24', '46'],
        ['Rating', '3/10', '2/10', '5/10', '8.5/10', '5/10', '2/10', '3/10', '6/10'],
    ]
)

add_heading('Annotated Scoring Matrix', level=1)

add_heading('Course Alignment — Fit with hypothesis tests, CIs, distributions, Bayesian methods', level=2)
add_table(
    ['Idea', 'Score', 'Justification'],
    [
        ['Weather-Tourism', '4', 'Correlation/regression possible but near-zero weather variation means non-significant tests. Time series better but beyond Sessions 1-5.'],
        ['Telecom Churn', '3', 'Natural framing is classification/prediction (ML, not stats thinking). Hypothesis tests can be forced in but feel unnatural.'],
        ['ESG Returns', '6', 'ANOVA across ESG quintiles, t-tests, F-test for volatility fit well. Needs multivariate regression to control confounders.'],
        ['Spotify', '9', 'Every method applies naturally: t-tests, ANOVA, CIs, chi-squared, non-parametric, Bayesian A/B test.'],
        ['Speed Dating', '6', 'Two-sample t-tests, paired t-tests, chi-squared, ANOVA all apply. 195 columns invite p-hacking.'],
        ['Diversification', '2', 'No clear hypothesis. Simulation-based, not inferential. Alpha needs factor models outside course scope.'],
        ['Dividends', '4', 'One clean two-sample t-test. ANOVA by yield quartile possible. Statistical work is thin.'],
        ['SG Pools', '7', 'Poisson goodness-of-fit (Session 3), chi-squared, Bayesian prior/posterior. Sparse data may limit power.'],
    ]
)

add_heading('Data Availability — Downloadable, clean, large enough, no scraping/merges', level=2)
add_table(
    ['Idea', 'Score', 'Justification'],
    [
        ['Weather-Tourism', '5', 'Three CSVs from data.gov.sg. Tiny (~526 rows), require merging, tourist arrivals end in 2015.'],
        ['Telecom Churn', '8', 'One-click Kaggle download, ~7K rows, clean, no missing. But synthetic IBM demo data.'],
        ['ESG Returns', '5', 'ESG CSV (503 rows) + stock prices available. Merging snapshot with daily time series is non-trivial.'],
        ['Spotify', '10', 'Single CSV, 114K rows x 20 cols, clean, no merge. Largest and most analysis-ready dataset.'],
        ['Speed Dating', '5', 'One CSV, 8.4K rows. But 195 columns with 65%+ missing requires extensive cleaning.'],
        ['Diversification', '1', 'No dataset identified. Building portfolio data from scratch is a data engineering project.'],
        ['Dividends', '4', 'Stock prices available. Dividend classification data not identified — critical gap.'],
        ['SG Pools', '3', 'Needs web scraping (SharePoint + dynamic loading). Census needs geo-matching. ~300 rows.'],
    ]
)

add_heading('Novelty — New perspective, counter-intuitive hypothesis, original angle', level=2)
add_table(
    ['Idea', 'Score', 'Justification'],
    [
        ['Weather-Tourism', '4', 'Weather-tourism studies exist. Singapore angle mild novelty. Equatorial constraint makes question uninteresting.'],
        ['Telecom Churn', '1', 'Least novel possible. 10,000+ Kaggle notebooks. Professor has certainly seen this before.'],
        ['ESG Returns', '4', 'Heavily researched (2,200+ studies). E/S/G decomposition adds some novelty but well-trodden.'],
        ['Spotify', '7', 'Counter-intuitive premise ("prefer sad music?") engaging. Genre + temporal angles are original.'],
        ['Speed Dating', '3', 'Famous 2006 dataset (Fisman & Iyengar). Hundreds of analyses. Self vs other mildly novel.'],
        ['Diversification', '3', '70-year-old theory (Markowitz 1952). Overlapping-assets is a factoid, not a hypothesis.'],
        ['Dividends', '2', 'Most debated, most analyzed topic in retail investing. Very limited novelty.'],
        ['SG Pools', '10', 'Most novel by far. No IS630 team has done this. Locally relevant, quirky, engaging.'],
    ]
)

add_heading('Statistical Depth — Can 6 people each apply distinct methods?', level=2)
add_table(
    ['Idea', 'Score', 'Justification'],
    [
        ['Weather-Tourism', '3', 'After correlations + regression, little left. Tiny dataset. 6 distinct threads hard to fill.'],
        ['Telecom Churn', '3', 'Steps form a linear pipeline, not 6 independent analyses. Most are prerequisites for the next.'],
        ['ESG Returns', '6', 'Overall test, E/S/G decomposition, sector ANOVA, volatility, Bayesian = 5-6 threads. 503 rows limits sizes.'],
        ['Spotify', '9', '6+ independent hypotheses each on 114K rows. Natural parallelism across features and genres.'],
        ['Speed Dating', '6', 'Gender diffs, self vs other, match predictors, demographics = 5 threads. Missing data may collapse some.'],
        ['Diversification', '2', 'Simulation + correlation + variance curve. None are formal hypothesis tests. Hard for 6 people.'],
        ['Dividends', '3', 'One t-test, one ANOVA, one F-test = 3 tests. Half the team has no statistical deliverable.'],
        ['SG Pools', '5', 'Poisson, chi-squared, Bayesian, demographics = 4 threads. Geo mapping = 5th. Sixth stretched.'],
    ]
)

add_heading('"So What?" Factor — Does the conclusion matter?', level=2)
add_table(
    ['Idea', 'Score', 'Justification'],
    [
        ['Weather-Tourism', '3', 'STB already knows weather doesn\'t drive tourism. No actionable insight.'],
        ['Telecom Churn', '4', 'Retention relevant, but dataset is synthetic. Recommendations go to a fictional company.'],
        ['ESG Returns', '5', 'Relevant to investors. Null result has muted impact. E/S/G decomposition adds value if one matters.'],
        ['Spotify', '7', 'Implications for labels, curators, artists. Everyone listens to music — visceral "So what?"'],
        ['Speed Dating', '3', 'No decision-maker benefits. No policy implication. "Unserious" label reflects weakness.'],
        ['Diversification', '4', 'Conclusion is textbook knowledge. No new insight from re-deriving known theory.'],
        ['Dividends', '4', 'Answer depends on time period, already well-established. Limited marginal insight.'],
        ['SG Pools', '9', 'Debunks common SG superstition. SMU audience relates — many have bought TOTO.'],
    ]
)

add_heading('Presentation Appeal — Engage non-specialist audience in 15 min?', level=2)
add_table(
    ['Idea', 'Score', 'Justification'],
    [
        ['Weather-Tourism', '4', 'Bland charts. Equatorial flatness = unexciting visuals. No narrative arc.'],
        ['Telecom Churn', '3', 'Generic problem, no emotional hook. Audience has seen this in other courses.'],
        ['ESG Returns', '5', 'Topical, can spark debate. But financial charts are dry. Null result = anticlimactic.'],
        ['Spotify', '9', 'Embed real songs. Audience recognizes tracks, debates "sad" label, has personal stakes.'],
        ['Speed Dating', '8', '"What do men vs women care about?" sparks engagement. Fun topic helps energy.'],
        ['Diversification', '3', 'Frontier plots are finance-class material, not storytelling. Niche audience.'],
        ['Dividends', '3', 'Standard finance fare. No emotional hook, no surprise, no participation angle.'],
        ['SG Pools', '8', '"Who here buys TOTO at a lucky outlet?" Instant engagement. Maps are striking.'],
    ]
)

add_heading('Low Risk — Low probability of null results or data roadblocks', level=2)
add_table(
    ['Idea', 'Score', 'Justification'],
    [
        ['Weather-Tourism', '3', '2.5C range almost guarantees non-significance. No fallback if hypothesis fails.'],
        ['Telecom Churn', '6', 'Low data risk (clean synthetic). High risk of generic, unremarkable analysis.'],
        ['ESG Returns', '4', 'Likely null result. Merge complexity consumes weeks. Survivorship bias undermines claims.'],
        ['Spotify', '7', 'If main hypothesis fails, 10+ features and 125 genres to explore. Guaranteed clean data.'],
        ['Speed Dating', '5', '65% missing may eliminate planned analyses. Small sample (552) limits power.'],
        ['Diversification', '2', 'No dataset = can\'t start. Simulation may not satisfy rubric\'s hypothesis testing demand.'],
        ['Dividends', '4', 'Time period dependency flips results. Thin analysis has little fallback.'],
        ['SG Pools', '4', 'Scrape may fail. Predetermined "no lucky outlets" limits surprise value.'],
    ]
)

add_heading('Work Split: Spotify (Recommended)', level=1)
add_table(
    ['Member', 'Role', 'Key Deliverables'],
    [
        ['M1', 'Project Lead + Data Prep', 'Clean dataset, data dictionary, EDA summary'],
        ['M2', 'Valence-Popularity Analysis', 't-test, Mann-Whitney U, CIs, effect size, scatter plots'],
        ['M3', 'Genre Segmentation', 'ANOVA / Kruskal-Wallis across genres, post-hoc tests'],
        ['M4', 'Temporal Trends', 'Mean valence by decade, trend tests, popularity evolution'],
        ['M5', 'Multi-Feature Analysis', 'Regression, correlation matrix, feature importance'],
        ['M6', 'Bayesian + Synthesis', 'Bayesian A/B test, credible intervals, report editing'],
    ]
)

add_heading('Work Split: All Other Ideas', level=1)
add_para('Detailed work splits for all 8 ideas are in the companion file Project_Evaluation.md.', italic=True)

add_heading('Final Summary', level=1)
add_table(
    ['Rank', 'Idea', 'Score', 'One-Line Reason'],
    [
        ['1', 'Spotify', '8.5/10', 'Largest clean dataset, most hypotheses, best course fit'],
        ['2', 'SG Pools', '6/10', 'Most novel, but data scraping and sparsity are risks'],
        ['3', 'Speed Dating', '5/10', 'Rich data but 65% missing, old, "unserious"'],
        ['4', 'ESG Returns', '5/10', 'Good depth but tiny dataset (503 rows), complex merge'],
        ['5', 'Weather-Tourism', '3/10', 'Near-zero weather variation — no signal'],
        ['6', 'Dividends', '3/10', 'One main t-test, well-trodden'],
        ['7', 'Telecom Churn', '2/10', 'Most overused topic in data science'],
        ['8', 'Diversification', '2/10', 'No dataset, no testable hypothesis'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Go with Spotify.')
run.bold = True
run.font.size = Pt(14)
p.add_run(' It is the only idea that simultaneously satisfies all seven evaluation criteria at a high level.')

doc.save('/home/dmgadmin/SMU/IS630/Project_Evaluation.docx')
print('Saved Project_Evaluation.docx')
