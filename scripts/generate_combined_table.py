from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)

h = doc.add_heading('IS630 Project Ideas — Combined Evaluation Matrix', level=0)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

p = doc.add_paragraph()
run = p.add_run('Honest scores. ')
run.bold = True
p.add_run('Each cell shows the score (1-10) and a one-line justification. Spotify scores revised down from initial assessment for candour on novelty, depth, and "So what?" limitations.')
doc.add_paragraph()

IDEAS = [
    'Weather-\nTourism',
    'Telecom\nChurn',
    'ESG\nReturns',
    'Spotify\n(Rec.)',
    'Speed\nDating',
    'Diversi-\nfication',
    'Divi-\ndends',
    'SG\nPools',
]

CRITERIA = [
    'Course\nAlignment',
    'Data\nAvailability',
    'Novelty',
    'Statistical\nDepth',
    '"So What?"\nFactor',
    'Presentation\nAppeal',
    'Low Risk',
]

DATA = {
    'Course\nAlignment': [
        (4, 'Correlation/regression possible but 2.5C range means non-significant tests'),
        (3, 'Natural framing is ML classification, not statistical thinking'),
        (6, 'ANOVA, t-tests, F-test fit well; needs regression for confounders'),
        (8, 'T-tests, ANOVA, CIs, chi-squared, non-parametric all natural; Bayesian slightly forced'),
        (6, 'T-tests, paired t-tests, chi-squared, ANOVA apply; 195 cols invite p-hacking'),
        (2, 'No testable hypothesis; simulation-based, not inferential statistics'),
        (4, 'One clean t-test, one ANOVA possible then work runs out'),
        (7, 'Poisson (Session 3), chi-squared, Bayesian prior/posterior direct course fit'),
    ],
    'Data\nAvailability': [
        (5, 'Three gov.sg CSVs but tiny (~526 rows), arrivals end 2015'),
        (8, 'One-click Kaggle, ~7K rows, clean but synthetic IBM demo data'),
        (5, 'ESG CSV (503 rows) + stock prices; merging snapshot with daily series non-trivial'),
        (10, 'Single CSV, 114K rows x 20 cols, clean, no merge best dataset by far'),
        (5, 'One CSV, 8.4K rows; 195 cols with 65%+ missing needs heavy cleaning'),
        (1, 'No dataset identified; building portfolio data is a separate project'),
        (4, 'Stock prices available; dividend classification data not identified'),
        (3, 'Needs web scraping (SharePoint); census needs geo-matching; ~300 rows'),
    ],
    'Novelty': [
        (4, 'Weather-tourism studies exist; equatorial climate makes answer obvious'),
        (1, '10,000+ Kaggle notebooks; professor has certainly seen this before'),
        (4, '2,200+ prior studies; E/S/G decomposition adds some originality'),
        (5, '"Prefer sad music?" is engaging, but Spotify analysis very common on Kaggle too'),
        (3, 'Famous 2006 dataset (Fisman & Iyengar); hundreds of analyses exist'),
        (3, '70-year-old Markowitz theory; overlapping-assets is a factoid'),
        (2, 'Most debated topic in retail investing; Vanguard/Morningstar covered it'),
        (7, 'Statistical normalization novel, but "lucky outlet" myth well-covered in SG media'),
    ],
    'Statistical\nDepth': [
        (3, 'After correlations + regression, little left; 6 threads hard to fill'),
        (3, 'Steps form a linear pipeline, not 6 independent analyses'),
        (6, 'Overall test, E/S/G, sector ANOVA, volatility, Bayesian = 5-6 threads'),
        (7, '6 hypotheses but all variations of "does feature X predict popularity?"; p-hacking risk'),
        (6, 'Gender, self-vs-other, match predictors, demographics = 5 threads'),
        (2, 'Simulation + correlation + variance curve; none are hypothesis tests'),
        (3, 'One t-test, one ANOVA, one F-test = 3 tests; half the team idle'),
        (5, 'Poisson, chi-squared, Bayesian, demographics = 4-5 threads; 6th stretched'),
    ],
    '"So What?"\nFactor': [
        (3, 'STB already knows weather does not drive tourism; no actionable insight'),
        (4, 'Retention relevant, but recommendations go to a fictional company'),
        (5, 'Relevant to investors; null result has muted impact'),
        (5, 'Relatable but not genuinely actionable; labels do not use student t-tests'),
        (3, 'No decision-maker benefits; proposer called it "unserious"'),
        (4, 'Conclusion is textbook knowledge; no new insight'),
        (4, 'Answer depends on time period; already well-established'),
        (9, 'Debunks common SG superstition; every Singaporean has an opinion'),
    ],
    'Presentation\nAppeal': [
        (4, 'Flat equatorial data = bland charts; no narrative arc'),
        (3, 'Generic problem, no emotional hook; seen in other courses'),
        (5, 'Topical ESG debate; but financial charts dry; null result anticlimactic'),
        (9, 'Embed real songs; audience recognises tracks, debates "sad" label'),
        (8, '"What do men vs women care about?" inherently entertaining'),
        (3, 'Frontier plots are finance-class material, not storytelling'),
        (3, 'Standard finance charts; no emotional hook or surprise'),
        (8, '"Who buys TOTO at a lucky outlet?" instant engagement; maps striking'),
    ],
    'Low Risk': [
        (3, '2.5C range guarantees non-significance; no fallback'),
        (6, 'Clean synthetic data; but high risk of generic unremarkable output'),
        (4, 'Likely null result; merge complexity; survivorship bias'),
        (6, 'Many features as fallback, but core signal likely weak (low R-sq); fishing risk'),
        (5, '65% missing may eliminate threads; small sample (552) limits power'),
        (2, 'No dataset = cannot start; simulation may not satisfy rubric'),
        (4, 'Time period dependency flips results; thin analysis, little fallback'),
        (4, 'Scrape may fail; ~300 rows sparse; predetermined "no" limits surprise'),
    ],
}

TOTALS = []
for i in range(8):
    total = sum(DATA[c][i][0] for c in CRITERIA)
    TOTALS.append(total)

num_cols = 1 + 8
num_rows = 1 + 7 + 1

table = doc.add_table(rows=num_rows, cols=num_cols)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

def set_cell(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bg_color:
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): bg_color,
            qn('w:val'): 'clear',
        })
        shading.append(shd)

def score_color(score):
    if score >= 8: return 'C6EFCE'
    if score >= 6: return 'FFEB9C'
    if score >= 4: return 'FFF2CC'
    return 'FFC7CE'

set_cell(table.rows[0].cells[0], 'Criterion', bold=True, size=9, bg_color='D9E2F3')
for j, idea in enumerate(IDEAS):
    bg = 'C6EFCE' if j == 3 else 'D9E2F3'
    set_cell(table.rows[0].cells[j + 1], idea, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color=bg)

for r, criterion in enumerate(CRITERIA):
    row = table.rows[r + 1]
    set_cell(row.cells[0], criterion, bold=True, size=8, bg_color='F2F2F2')
    for j in range(8):
        score, justification = DATA[criterion][j]
        cell_text = f'{score} — {justification}'
        bg = score_color(score)
        if j == 3:
            bg = 'B7E1CD'
        set_cell(row.cells[j + 1], cell_text, size=7, bg_color=bg)

total_row = table.rows[num_rows - 1]
set_cell(total_row.cells[0], 'TOTAL\n(/70)', bold=True, size=9, bg_color='D9E2F3')

ratings = {
    0: '3/10', 1: '2/10', 2: '5/10', 3: '7/10',
    4: '5.5/10', 5: '2/10', 6: '3/10', 7: '6/10',
}
for j in range(8):
    t = TOTALS[j]
    rating = ratings.get(j, '')
    bg = 'B7E1CD' if j == 3 else score_color(t // 7)
    set_cell(total_row.cells[j + 1], f'{t}/70\n({rating})', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color=bg)

for row in table.rows:
    row.cells[0].width = Cm(2.2)
    for j in range(1, 9):
        row.cells[j].width = Cm(3.2)

doc.add_paragraph()

h2 = doc.add_heading('Honesty Note: Spotify Score Revision', level=1)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

revisions = doc.add_table(rows=6, cols=4)
revisions.style = 'Light Grid Accent 1'
for i, (crit, orig, revised, reason) in enumerate([
    ('Criterion', 'Original', 'Revised', 'What was inflated'),
    ('Novelty', '7', '5', 'Spotify audio feature analysis is nearly as common on Kaggle as churn'),
    ('Statistical Depth', '9', '7', 'The 6 hypotheses are all variations of "does feature X predict popularity?"'),
    ('"So What?"', '7', '5', 'Relatable but not genuinely actionable; no real decision-maker uses this'),
    ('Low Risk', '7', '6', '10+ features is also 10+ fishing opportunities; core signal likely weak'),
    ('Course Alignment', '9', '8', 'Bayesian angle slightly forced; rest genuinely fits'),
]):
    for j, val in enumerate([crit, orig, revised, reason]):
        cell = revisions.rows[i].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Verdict: ')
run.bold = True
run.font.size = Pt(11)
p.add_run('Spotify remains #1 (50/70). SG Pools drops to 43/70 after novelty revision — the "lucky outlet" myth is well-covered in Singaporean media (TheSmartLocal, Goody Feed, Mothership all rank outlets by raw wins). The novel angle (normalizing by estimated ticket volume via HDB density) hasn\'t been done, but the qualitative answer is already known. Gap to Spotify is now 7 points.')

doc.add_paragraph()
h3 = doc.add_heading('Final Ranking (Honest Scores)', level=1)
for run in h3.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

final = doc.add_table(rows=9, cols=3)
final.style = 'Light Grid Accent 1'
for i, (rank, idea, score) in enumerate([
    ('Rank', 'Idea', 'Honest Score'),
    ('1', 'Spotify', '50/70 (7.0/10)'),
    ('2', 'SG Pools', '43/70 (6.0/10)'),
    ('3', 'Speed Dating', '38/70 (5.5/10)'),
    ('4', 'ESG Returns', '35/70 (5.0/10)'),
    ('5', 'Weather-Tourism', '26/70 (3.0/10)'),
    ('6', 'Dividends', '24/70 (3.0/10)'),
    ('7', 'Telecom Churn', '28/70 (2.0/10)'),
    ('8', 'Diversification', '17/70 (2.0/10)'),
]):
    for j, val in enumerate([rank, idea, score]):
        cell = final.rows[i].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.save('/home/dmgadmin/SMU/IS630/Combined_Table.docx')
print('Saved Combined_Table.docx')
