#!/usr/bin/env python3
"""Convert TOTO_Project_Writeup.md to a formatted Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

def parse_markdown_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and not all(c in '|-: ' for c in line):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
    return rows

def add_table_from_rows(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < len(table.columns):
                cell = table.cell(i, j)
                cell.text = cell_text.replace('**', '').replace('`', '')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    return table

def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.color.rgb = RGBColor(0, 51, 102)

    with open(md_path, 'r') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        if line.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
                p.paragraph_format.left_indent = Inches(0.3)
                code_lines = []
                in_code_block = False
            else:
                if in_table and table_lines:
                    rows = parse_markdown_table(table_lines)
                    add_table_from_rows(doc, rows)
                    table_lines = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            rows = parse_markdown_table(table_lines)
            add_table_from_rows(doc, rows)
            table_lines = []
            in_table = False

        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            doc.add_heading(text, level=1)
        elif line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
        elif line.startswith('---'):
            pass
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = re.sub(r'^[\s]*[-*]\s+', '', line)
            text = text.replace('**', '').replace('`', '').replace('*', '')
            doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            text = text.replace('**', '').replace('`', '').replace('*', '')
            doc.add_paragraph(text, style='List Number')
        elif line.strip() == '':
            pass
        else:
            text = line.strip()
            if text.startswith('**') and text.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(text.replace('**', ''))
                run.bold = True
            else:
                text = text.replace('**', '').replace('`', '')
                if text.startswith('*') and text.endswith('*'):
                    p = doc.add_paragraph()
                    run = p.add_run(text.strip('*'))
                    run.italic = True
                else:
                    doc.add_paragraph(text)

        i += 1

    if in_table and table_lines:
        rows = parse_markdown_table(table_lines)
        add_table_from_rows(doc, rows)

    doc.save(docx_path)
    print(f"Saved to {docx_path}")

if __name__ == '__main__':
    convert_md_to_docx(
        '/home/dmgadmin/SMU/IS630/TOTO_Project_Writeup.md',
        '/home/dmgadmin/SMU/IS630/TOTO_Project_Writeup.docx'
    )
